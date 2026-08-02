from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "AfranHospital_LAN_Deployment_Guide.docx"
BLUE = "2E74B5"
NAVY = "0B2545"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "EAF5EE"
GOLD = "FFF8E8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color="000000", bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", style=None, size=11, color="000000", bold=False, italic=False, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        set_font(p.add_run(text), size, color, bold, italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(label + "  "), 10.5, NAVY, True)
    set_font(p.add_run(text), 10.5, NAVY)
    add_para(doc, after=4)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), 10, NAVY, True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value), 10)
    add_para(doc, after=8)
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def main():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("AFRAN HOSPITAL  |  LAN DEPLOYMENT GUIDE"), 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Afran Hospital Queue System  •  Windows 10/11  •  Self-contained deployment"), 8.5, MUTED)

    add_para(doc, "AFRAN HOSPITAL", size=12, color=BLUE, bold=True, after=2)
    add_para(doc, "Queue System LAN Deployment Guide", size=27, color=NAVY, bold=True, after=4)
    add_para(doc, "Installation, configuration, and daily operating instructions", size=13, color=MUTED, after=18)
    add_callout(doc, "Quick start", "Copy the release folder to the server and client computers. Start the Queue API on the server, then point each client launcher to the server IP address.", GREEN)

    add_para(doc, "Document purpose", style="Heading 1")
    add_para(doc, "This guide explains how to deploy the Afran Hospital Queue System across a hospital local area network (LAN). The system includes a patient registration kiosk, a doctor station, a waiting-room display, and a Queue API backed by SQLite.")
    add_para(doc, "The launcher is the single user interface for starting and monitoring the modules. In a multi-computer deployment, the same launcher is used on every computer, but each computer starts only its assigned module.")

    add_para(doc, "Deployment layout", style="Heading 1")
    add_table(doc, ["Computer", "Run", "Purpose"], [
        ("Server", "Queue API", "Stores tickets and provides the LAN service on port 5000."),
        ("Kiosk / Reception", "Patient Kiosk", "Registers patients and prints queue tickets."),
        ("Doctor desk", "Doctor Station", "Calls, recalls, and completes tickets."),
        ("Waiting room", "Waiting Room TV", "Shows the current ticket and waiting queue."),
    ], [1800, 2400, 5160])

    add_para(doc, "Before you begin", style="Heading 1")
    for text in [
        "Use Windows 10 or Windows 11, 64-bit, on each computer.",
        "Connect all computers to the same trusted hospital LAN.",
        "Use a fixed or reserved IP address for the server computer.",
        "Have a thermal or Windows-supported printer connected to the kiosk computer if tickets must be printed.",
        "Copy the complete release folder generated by deploy-now.ps1; do not copy only the executable.",
    ]:
        add_bullet(doc, text)

    add_para(doc, "1. Install and start the server", style="Heading 1")
    add_number(doc, "Copy the complete release folder to the server computer, for example C:\\AfranHospital\\release.")
    add_number(doc, "Open release\\Start-AfranHospital.bat.")
    add_number(doc, "In the launcher, click Start API and wait until the status shows Online.")
    add_number(doc, "Open PowerShell and run ipconfig. Record the server IPv4 address, such as 192.168.1.10.")
    add_number(doc, "Keep the Queue API running while the kiosk, doctor, and TV computers are in use.")
    add_callout(doc, "Important", "Start the Queue API on the server only. Client computers should point to the server and should not start their own API instance.", GOLD)

    add_para(doc, "2. Configure Windows Firewall", style="Heading 1")
    add_para(doc, "On the server computer, open PowerShell as Administrator and allow inbound TCP traffic on port 5000:")
    p = add_para(doc, after=8)
    p.paragraph_format.left_indent = Inches(0.25)
    set_font(p.add_run('netsh advfirewall firewall add rule name="Afran Queue API" dir=in action=allow protocol=TCP localport=5000'), 10, NAVY, True)
    add_para(doc, "If the hospital firewall is centrally managed, ask the network administrator to allow TCP 5000 from the hospital LAN to the server IP address.")

    add_para(doc, "3. Install and configure client computers", style="Heading 1")
    add_number(doc, "Copy the complete release folder to the kiosk, doctor, and TV computers.")
    add_number(doc, "Open release\\Start-AfranHospital.bat on the client computer.")
    add_number(doc, "In the API address field, replace http://localhost:5000 with the server address, for example http://192.168.1.10:5000.")
    add_number(doc, "Confirm the API status changes to Online.")
    add_number(doc, "Launch only the module assigned to that computer.")
    add_table(doc, ["Computer", "Launcher action"], [
        ("Kiosk", "Launch Kiosk"),
        ("Doctor desk", "Launch Doctor"),
        ("Waiting room TV", "Launch TV"),
    ], [3000, 6360])
    add_callout(doc, "Do not", "Do not click Start API on a client computer. Multiple API instances create separate databases and can split the queue.", GOLD)

    add_para(doc, "4. Test the LAN connection", style="Heading 1")
    add_para(doc, "From each client computer, replace the sample IP address with the actual server address and run:")
    p = add_para(doc, after=8)
    p.paragraph_format.left_indent = Inches(0.25)
    set_font(p.add_run("Test-NetConnection 192.168.1.10 -Port 5000"), 10, NAVY, True)
    add_para(doc, "A successful test should show TcpTestSucceeded: True. You can also open this address in a browser to confirm the API is running:")
    p = add_para(doc, after=8)
    p.paragraph_format.left_indent = Inches(0.25)
    set_font(p.add_run("http://192.168.1.10:5000/"), 10, NAVY, True)
    add_para(doc, "The response should indicate that the Afran Queue API is Running.")

    add_para(doc, "5. Daily operating procedure", style="Heading 1")
    for text in [
        "Start the server computer and open the launcher.",
        "Start the Queue API and confirm the Online status.",
        "Start the Kiosk, Doctor, and TV modules on their assigned computers.",
        "Patients select their language and gender at the Kiosk and collect a printed ticket.",
        "The doctor uses Doctor Station to call the next ticket, recall it, or complete it.",
        "The TV displays the now-serving ticket and the next waiting tickets.",
    ]:
        add_number(doc, text)

    add_para(doc, "6. Data and backup", style="Heading 1")
    add_para(doc, "The Queue API creates its SQLite database automatically in the API folder:")
    p = add_para(doc, after=8)
    p.paragraph_format.left_indent = Inches(0.25)
    set_font(p.add_run("release\\QueueApi\\SQLite.db"), 10, NAVY, True)
    add_para(doc, "Back up this file when the Queue API is stopped. Keep daily backups on a separate drive or approved hospital storage. Do not delete the file unless you intentionally want to reset the queue history.")

    add_para(doc, "Troubleshooting", style="Heading 1")
    add_table(doc, ["Problem", "Check"], [
        ("Client shows Offline", "Confirm the API address uses the server IP, the API is running, and port 5000 is allowed through the server firewall."),
        ("TcpTestSucceeded is False", "Check LAN connectivity, server IP address, network profile, and Windows Firewall."),
        ("Tickets do not synchronize", "Ensure every client uses the same server URL and that only one Queue API is running."),
        ("Kiosk does not print", "Set a default Windows printer and test printing from Windows. The queue still works if printing is unavailable."),
        ("TV has no current ticket", "Confirm the Doctor Station is connected to the same API and has called a waiting ticket."),
    ], [2700, 6660])

    add_para(doc, "Useful command-line modes", style="Heading 1")
    add_para(doc, "The launcher is recommended for deployment. Direct modes remain available for troubleshooting or dedicated shortcuts:")
    add_table(doc, ["Command", "Opens"], [
        ("AfranHospitalKiosk.exe", "Unified launcher"),
        ("AfranHospitalKiosk.exe kiosk", "Patient Kiosk"),
        ("AfranHospitalKiosk.exe doctor", "Doctor Station"),
        ("AfranHospitalKiosk.exe tv", "Waiting Room TV"),
    ], [4200, 5160])
    add_callout(doc, "Support checklist", "When requesting support, provide the server IP address, the affected computer role, the launcher API address, and the result of Test-NetConnection on port 5000.", LIGHT_GRAY)

    doc.core_properties.title = "Afran Hospital Queue System LAN Deployment Guide"
    doc.core_properties.subject = "Deployment and daily use instructions"
    doc.core_properties.author = "Afran Hospital"
    doc.core_properties.keywords = "Afran Hospital, queue system, LAN, deployment"
    doc.save(OUT)


if __name__ == "__main__":
    main()
